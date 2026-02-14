"""
Premium CV Export Service — ATS-Optimized DOCX Generator.

Produces clean, B&W, ATS-parser-friendly Word documents from the
enhanced CV text.  No tables, no graphics, no columns — just clean
section headers and body text that every ATS can parse.
"""

from __future__ import annotations
import io
import json
import re
from html import unescape
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Section patterns we recognise in the enhanced CV text ─────────────
_SECTION_HEADERS = [
    "PROFESSIONAL SUMMARY", "SUMMARY", "PROFILE", "OBJECTIVE",
    "EXPERIENCE", "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE",
    "EMPLOYMENT HISTORY", "WORK HISTORY",
    "SKILLS", "TECHNICAL SKILLS", "CORE COMPETENCIES", "KEY SKILLS",
    "EDUCATION", "ACADEMIC BACKGROUND",
    "CERTIFICATIONS", "CERTIFICATES", "LICENSES",
    "PROJECTS", "KEY PROJECTS",
    "ACHIEVEMENTS", "ACCOMPLISHMENTS",
    "LANGUAGES", "ADDITIONAL INFORMATION", "REFERENCES",
]

_SECTION_RE = re.compile(
    r"^(?:#{1,3}\s+)?("
    + "|".join(re.escape(h) for h in _SECTION_HEADERS)
    + r")\s*:?\s*$",
    re.IGNORECASE | re.MULTILINE,
)


def _extract_enhanced_cv_from_json_like(text: str) -> str:
    """Extract enhanced_cv from JSON-like payloads, fallback to input text."""
    if not text:
        return ""
    raw = text.strip()
    try:
        payload = json.loads(raw)
        if isinstance(payload, dict):
            cv = payload.get("enhanced_cv")
            if isinstance(cv, str):
                return cv
    except Exception:
        pass

    # Tolerant extraction for malformed/truncated JSON.
    key_match = re.search(r'"enhanced_cv"\s*:\s*"', raw)
    if key_match:
        chars: list[str] = []
        escaped = False
        idx = key_match.end()
        while idx < len(raw):
            ch = raw[idx]
            if escaped:
                chars.append(ch)
                escaped = False
                idx += 1
                continue
            if ch == "\\":
                chars.append(ch)
                escaped = True
                idx += 1
                continue
            if ch == '"':
                break
            chars.append(ch)
            idx += 1

        extracted = "".join(chars).strip()
        if extracted:
            try:
                return json.loads(f'"{extracted}"')
            except Exception:
                return (
                    extracted.replace('\\"', '"')
                    .replace("\\n", "\n")
                    .replace("\\r", "\n")
                    .replace("\\t", "\t")
                    .replace("\\/", "/")
                    .replace("\\\\", "\\")
                    .strip()
                )

    return raw


def _html_to_plain_text(text: str) -> str:
    """Convert simple resume HTML to ATS-safe plain text."""
    if not text:
        return ""
    out = text
    # Normalize fenced blocks, if any.
    out = re.sub(r"^```(?:html|json)?\s*", "", out, flags=re.IGNORECASE).strip()
    out = re.sub(r"\s*```$", "", out).strip()
    # Convert likely JSON payloads.
    out = _extract_enhanced_cv_from_json_like(out)
    # Normalize common block tags to new lines before stripping tags.
    out = re.sub(r"<br\s*/?>", "\n", out, flags=re.IGNORECASE)
    out = re.sub(r"</(p|div|h1|h2|h3|h4|h5|h6|section|article)>", "\n", out, flags=re.IGNORECASE)
    out = re.sub(r"<li[^>]*>", "- ", out, flags=re.IGNORECASE)
    out = re.sub(r"</li>", "\n", out, flags=re.IGNORECASE)
    out = re.sub(r"</(ul|ol)>", "\n", out, flags=re.IGNORECASE)
    # Drop all remaining tags.
    out = re.sub(r"<[^>]+>", "", out)
    # Decode HTML entities + normalize escaped newlines that may come from JSON strings.
    out = unescape(out).replace("\\n", "\n").replace("\\t", "\t")
    # Cleanup spacing.
    out = re.sub(r"\r\n?", "\n", out)
    out = re.sub(r"\n{3,}", "\n\n", out)
    out = re.sub(r"[ \t]{2,}", " ", out)
    return out.strip()


def _validate_enhancement_for_export(enhanced_content: dict) -> tuple[bool, list[str]]:
    """
    Validates enhanced CV content before export to DOCX/HTML.
    Returns (is_valid, list_of_errors).
    """
    errors: list[str] = []
    payload = enhanced_content if isinstance(enhanced_content, dict) else {}
    content_text = str(payload).lower()

    # Gate 1: Required sections present
    required_sections = [
        "professional_summary",
        "core_competencies",
        "professional_experience",
        "education",
        "contact_information",
    ]
    section_aliases = {
        "professional_summary": ["professional summary", "summary", "profile", "objective"],
        "core_competencies": ["core competencies", "skills", "technical skills", "key skills"],
        "professional_experience": ["professional experience", "experience", "work experience", "employment history"],
        "education": ["education", "academic background"],
        "contact_information": ["contact", "email", "phone", "linkedin"],
    }
    for section in required_sections:
        aliases = [section, section.replace("_", " ")] + section_aliases.get(section, [])
        if not any(alias in content_text for alias in aliases):
            errors.append(f"Required section missing: {section}")

    # Gate 2: No JSON-wrapper leakage
    json_leak_patterns = [
        "```json",
        "```",
        '{"',
        '"}',
        "\\n\\n",
        "\\u00",
        '"sections":',
        '"content":',
    ]
    rendered = payload.get("rendered_text", "") or payload.get("content", "")
    rendered_text = str(rendered)
    for pattern in json_leak_patterns:
        if pattern in rendered_text:
            errors.append(f"JSON/code leakage detected: '{pattern}' found in rendered output")

    # Gate 3: Semantic HTML validation (for HTML exports)
    if str(payload.get("format", "")).lower() == "html":
        html_content = str(payload.get("html", "") or "")
        if html_content:
            html_lower = html_content.lower()
            if "<h1" not in html_lower and "<h2" not in html_lower:
                errors.append("HTML export lacks heading structure (no h1/h2 tags)")
            if html_content.count("<") < 5:
                errors.append("HTML export appears to be plain text, not semantic HTML")
            for tag in ["div", "section", "ul", "ol", "table"]:
                opens = html_lower.count(f"<{tag}")
                closes = html_lower.count(f"</{tag}>")
                if opens != closes:
                    errors.append(f"HTML validation: mismatched <{tag}> tags ({opens} open, {closes} close)")

    return len(errors) == 0, errors


def generate_premium_docx(
    enhanced_cv_text: str,
    job_title: str,
    company: str,
    candidate_name: str | None = None,
    skills_matched: list[str] | None = None,
    skills_missing: list[str] | None = None,
) -> bytes:
    """Return in-memory DOCX bytes for the ATS-optimised premium CV.

    Rules for ATS friendliness:
    • Single-column layout
    • No tables, text boxes, or floating images
    • Standard fonts (Calibri / Arial)
    • Section headings as styled Heading 1
    • Bullet points as plain list paragraphs
    """
    enhanced_cv_text = _html_to_plain_text(enhanced_cv_text)
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ── Default font ──────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    # ── Heading styles ────────────────────────────────────────────────
    h1_style = doc.styles["Heading 1"]
    h1_style.font.name = "Calibri"
    h1_style.font.size = Pt(13)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h1_style.paragraph_format.space_before = Pt(14)
    h1_style.paragraph_format.space_after = Pt(4)
    h1_style.paragraph_format.keep_with_next = True

    h2_style = doc.styles["Heading 2"]
    h2_style.font.name = "Calibri"
    h2_style.font.size = Pt(11)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x2A, 0x2A, 0x2A)
    h2_style.paragraph_format.space_before = Pt(10)
    h2_style.paragraph_format.space_after = Pt(2)

    # ── Candidate name (top) ──────────────────────────────────────────
    name = candidate_name or _extract_name(enhanced_cv_text) or "Candidate"
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(name.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # ── Targeted position line ────────────────────────────────────────
    target_para = doc.add_paragraph()
    target_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = target_para.add_run(f"Targeting: {job_title} at {company}")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── Thin separator ────────────────────────────────────────────────
    _add_separator(doc)

    # ── Parse enhanced text into sections ─────────────────────────────
    sections = _split_into_sections(enhanced_cv_text)

    for section_title, section_body in sections:
        if section_title:
            doc.add_heading(section_title.title(), level=1)

        for line in section_body.strip().splitlines():
            stripped = line.strip()
            if not stripped:
                continue

            # Sub-header detection (e.g. job title / company / dates)
            if _is_subheader(stripped):
                doc.add_heading(stripped, level=2)
            # Bullet point
            elif stripped.startswith(("•", "-", "–", "▪", "►", "*")):
                bullet_text = stripped.lstrip("•-–▪►* ").strip()
                p = doc.add_paragraph(bullet_text, style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
            else:
                doc.add_paragraph(stripped)

    # ── Skills keyword footer (ATS keyword injection) ─────────────────
    if skills_matched or skills_missing:
        _add_separator(doc)
        doc.add_heading("Core Competencies", level=1)
        all_skills = list(skills_matched or [])
        # Include some missing skills too — the CV was enhanced to cover them
        kw_text = " | ".join(all_skills[:30])
        if kw_text:
            p = doc.add_paragraph(kw_text)
            p.runs[0].font.size = Pt(9.5) if p.runs else None

    # ── Footer with generation date ───────────────────────────────────
    _add_separator(doc)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d')} — "
        f"ATS-optimized for {job_title} at {company}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True

    # ── Serialise ─────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Helpers ────────────────────────────────────────────────────────────

def _add_separator(doc: Document):
    """Add a thin horizontal rule via a paragraph border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("─" * 72)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def _extract_name(text: str) -> str | None:
    """Try to pull the candidate name from the first non-empty line."""
    for line in text.splitlines()[:5]:
        stripped = line.strip().strip("#").strip()
        if stripped and len(stripped) < 60 and not any(c in stripped for c in ":@|"):
            return stripped
    return None


def _split_into_sections(text: str) -> list[tuple[str | None, str]]:
    """Split the CV text by recognised section headers."""
    parts: list[tuple[str | None, str]] = []
    positions = [(m.start(), m.group(1)) for m in _SECTION_RE.finditer(text)]

    if not positions:
        return [(None, text)]

    # Content before first header
    if positions[0][0] > 0:
        parts.append((None, text[: positions[0][0]]))

    for i, (pos, header) in enumerate(positions):
        end = positions[i + 1][0] if i + 1 < len(positions) else len(text)
        body = text[pos + len(header) :end].strip().lstrip(":").strip()
        parts.append((header, body))

    return parts


def _is_subheader(line: str) -> bool:
    """Heuristic: short line in title-case or ALL CAPS, no punctuation."""
    if len(line) > 80:
        return False
    if line.endswith((".","!","?",":")):
        return False
    # Lines like "Senior Data Engineer | Acme Corp | Jan 2020 – Present"
    if "|" in line and "–" in line:
        return True
    # ALL CAPS short line that isn't a known section
    if line.isupper() and len(line.split()) <= 8:
        return True
    return False
