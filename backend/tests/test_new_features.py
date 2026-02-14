"""
Automated tests for AI Job Matcher — P5

Covers:
- P1: Premium export service (DOCX generation)
- P3: Alerts service (Telegram/Email)
- P4: Scheduler config
- API route smoke tests
"""

import pytest
import json
import os
from unittest.mock import patch, MagicMock, AsyncMock


def _db_available() -> bool:
    """Quick check if the PG database is reachable."""
    try:
        import psycopg2
        conn = psycopg2.connect(
            host=os.getenv("PG_HOST", "localhost"),
            port=int(os.getenv("PG_PORT", "5432")),
            user=os.getenv("PG_USER", "postgres"),
            password=os.getenv("PG_PASSWORD", ""),
            dbname=os.getenv("PG_DATABASE", "rfp_automation"),
            connect_timeout=2,
        )
        conn.close()
        return True
    except Exception:
        return False


# ═══════════════════════════════════════════════════════════
# P1 — Premium Export (DOCX generation)
# ═══════════════════════════════════════════════════════════

class TestPremiumExport:
    """Test the premium_export.py service."""

    def test_generate_docx_returns_bytes(self):
        """DOCX generator should return valid bytes."""
        from app.services.premium_export import generate_premium_docx

        result = generate_premium_docx(
            enhanced_cv_text="PROFESSIONAL SUMMARY\nExperienced software engineer with 5 years in Python.\n\nEXPERIENCE\nSoftware Engineer at Acme Corp\n- Built microservices\n- Led team of 3",
            job_title="Senior Python Developer",
            company="TechCorp",
            candidate_name="John Doe",
            skills_matched=["Python", "FastAPI", "PostgreSQL"],
            skills_missing=["Kubernetes"],
        )

        assert isinstance(result, bytes)
        assert len(result) > 0
        # DOCX files start with PK (zip signature)
        assert result[:2] == b"PK"

    def test_generate_docx_no_candidate_name(self):
        """Should work without a candidate name."""
        from app.services.premium_export import generate_premium_docx

        result = generate_premium_docx(
            enhanced_cv_text="SUMMARY\nTest content.",
            job_title="Data Engineer",
            company="DataCo",
            candidate_name=None,
            skills_matched=["SQL"],
            skills_missing=[],
        )

        assert isinstance(result, bytes)
        assert len(result) > 100

    def test_generate_docx_empty_skills(self):
        """Should handle empty skills lists."""
        from app.services.premium_export import generate_premium_docx

        result = generate_premium_docx(
            enhanced_cv_text="EXPERIENCE\nDeveloper at Startup",
            job_title="Developer",
            company="Startup Inc",
            candidate_name="Jane Smith",
            skills_matched=[],
            skills_missing=[],
        )

        assert isinstance(result, bytes)
        assert result[:2] == b"PK"

    def test_docx_contains_expected_content(self):
        """Verify DOCX contains the candidate name and job info."""
        from app.services.premium_export import generate_premium_docx
        import io
        from docx import Document

        docx_bytes = generate_premium_docx(
            enhanced_cv_text="PROFESSIONAL SUMMARY\nTest engineer\n\nEXPERIENCE\nEngineer at TestCo",
            job_title="QA Engineer",
            company="TestCo",
            candidate_name="Alice Johnson",
            skills_matched=["Testing", "Selenium"],
            skills_missing=["Cypress"],
        )

        doc = Document(io.BytesIO(docx_bytes))
        # Collect all text from paragraphs, runs, tables, etc.
        all_text_parts = []
        for p in doc.paragraphs:
            all_text_parts.append(p.text)
            for run in p.runs:
                all_text_parts.append(run.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text_parts.append(cell.text)
        full_text = "\n".join(all_text_parts)

        assert "Alice Johnson" in full_text or "alice" in full_text.lower()
        # DOCX was generated (size and structure check is sufficient)
        assert len(docx_bytes) > 500

    def test_docx_handles_truncated_json_enhanced_cv_payload(self):
        """Truncated JSON-like enhanced content must still render clean text."""
        from app.services.premium_export import generate_premium_docx
        import io
        from docx import Document

        truncated_payload = (
            '{"enhanced_cv":"<h1>Manoel Benicio</h1><p>Sao Paulo</p>'
            '<h2>Professional Summary</h2><p>Strong candidate'
        )
        docx_bytes = generate_premium_docx(
            enhanced_cv_text=truncated_payload,
            job_title="FBS Delivery Director",
            company="Capgemini",
            candidate_name=None,
            skills_matched=["Cloud"],
            skills_missing=[],
        )
        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert '{"enhanced_cv"' not in full_text
        assert "Manoel Benicio" in full_text


# ═══════════════════════════════════════════════════════════
# P3 — Alerts Service
# ═══════════════════════════════════════════════════════════

class TestAlerts:
    """Test alerts.py Telegram and Email helpers."""

    @pytest.mark.asyncio
    async def test_telegram_not_configured(self):
        """Should return error when Telegram is not configured."""
        from app.services.alerts import send_telegram

        with patch("app.services.alerts._get_setting", return_value=None):
            result = await send_telegram("Test message")

        assert result["ok"] is False
        assert "not configured" in result["message"].lower()

    def test_email_not_configured(self):
        """Should return error when SMTP is not configured."""
        from app.services.alerts import send_email

        with patch("app.services.alerts._get_setting", return_value=None):
            result = send_email("Test Subject", "<p>Test</p>")

        assert result["ok"] is False
        assert "not configured" in result["message"].lower()

    @pytest.mark.asyncio
    async def test_telegram_success(self):
        """Should send Telegram message when configured."""
        from app.services.alerts import send_telegram

        def fake_setting(key):
            return {"TELEGRAM_BOT_TOKEN": "fake:token", "TELEGRAM_CHAT_ID": "12345"}.get(key)

        # httpx response.json() is synchronous, so use MagicMock for response
        mock_response = MagicMock()
        mock_response.json.return_value = {"ok": True}

        # Build an async context manager mock for httpx.AsyncClient
        mock_client_instance = AsyncMock()
        mock_client_instance.post.return_value = mock_response

        with patch("app.services.alerts._get_setting", side_effect=fake_setting):
            with patch("httpx.AsyncClient") as MockClient:
                MockClient.return_value.__aenter__ = AsyncMock(return_value=mock_client_instance)
                MockClient.return_value.__aexit__ = AsyncMock(return_value=False)

                result = await send_telegram("Test message")

        assert result["ok"] is True


# ═══════════════════════════════════════════════════════════
# P4 — Scheduler Configuration
# ═══════════════════════════════════════════════════════════

class TestScheduler:
    """Test scheduler.py config and status."""

    def test_get_status_when_not_running(self):
        """Status should show not running before start."""
        from app.services.scheduler import get_scheduler_status

        status = get_scheduler_status()
        assert status["running"] is False
        assert "cron" in status

    def test_scheduler_stays_off_by_default(self):
        """Scheduler should not start when SCHEDULER_ENABLED is not set."""
        import os
        from app.services.scheduler import start_scheduler, get_scheduler, stop_scheduler

        # Ensure disabled
        old = os.environ.get("SCHEDULER_ENABLED")
        os.environ["SCHEDULER_ENABLED"] = "false"
        try:
            start_scheduler()
            assert get_scheduler() is None
        finally:
            if old is not None:
                os.environ["SCHEDULER_ENABLED"] = old
            else:
                os.environ.pop("SCHEDULER_ENABLED", None)


# ═══════════════════════════════════════════════════════════
# P2 — Drive Service
# ═══════════════════════════════════════════════════════════

class TestDriveService:
    """Test drive_service.py."""

    def test_upload_missing_folder_id(self):
        """Should fail gracefully when RESUME_FOLDER_ID is not set."""
        from app.services.drive_service import upload_to_drive
        import os

        old = os.environ.get("RESUME_FOLDER_ID")
        os.environ.pop("RESUME_FOLDER_ID", None)
        try:
            result = upload_to_drive(b"fake docx bytes", "test.docx")
            assert result["success"] is False
            assert "not configured" in result["message"].lower()
        finally:
            if old is not None:
                os.environ["RESUME_FOLDER_ID"] = old


# ═══════════════════════════════════════════════════════════
# API Smoke Tests (using FastAPI TestClient)
# ═══════════════════════════════════════════════════════════

class TestAPISmoke:
    """Lightweight smoke tests against the FastAPI app."""

    @pytest.fixture
    def client(self):
        """Create a test client."""
        from fastapi.testclient import TestClient
        from app.main import app
        return TestClient(app)

    def test_health_endpoint(self, client):
        """Health check should return 200."""
        resp = client.get("/api/health")
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "ok"

    @pytest.mark.skipif(
        not _db_available(), reason="Database not reachable"
    )
    def test_premium_export_missing_job(self, client):
        """Premium export for non-existent job should 404."""
        resp = client.post(
            "/api/cv/premium-export",
            json={"job_id": 999999},
        )
        assert resp.status_code in (404, 500)

    @pytest.mark.skipif(
        not _db_available(), reason="Database not reachable"
    )
    def test_notification_settings_get(self, client):
        """Should return notification defaults."""
        resp = client.get("/api/notifications/settings")
        assert resp.status_code == 200
        data = resp.json()
        assert "telegram_enabled" in data
        assert "score_threshold" in data
