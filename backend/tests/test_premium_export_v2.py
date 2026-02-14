import io
import re

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _render_semantic_doc(html: str) -> Document:
    from app.services import premium_export as pe

    doc = Document()
    if hasattr(pe, "_walk_html_to_docx"):
        pe._walk_html_to_docx(doc, html)
        return doc

    walker_cls = getattr(pe, "_DocxHtmlWalker", None)
    if walker_cls is not None:
        walker = walker_cls(doc)
        walker.feed(html)
        if hasattr(walker, "close"):
            walker.close()
        return doc

    pytest.fail("Semantic HTML->DOCX converter not implemented (_walk_html_to_docx or _DocxHtmlWalker missing)")


def _all_doc_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


# Module E: Semantic HTML->DOCX Converter (E1-E10)

def test_h1_becomes_centered_title():
    doc = _render_semantic_doc("<h1>Manoel Benicio</h1>")

    assert doc.paragraphs
    p = doc.paragraphs[0]
    assert "Manoel Benicio" in p.text
    assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert any(run.bold for run in p.runs)


def test_h2_becomes_heading1():
    doc = _render_semantic_doc("<h2>Experience</h2>")

    p = doc.paragraphs[0]
    assert "Experience" in p.text
    assert "heading" in p.style.name.lower()


def test_h3_becomes_heading2():
    doc = _render_semantic_doc("<h3>Head Strategic Business Development</h3>")

    p = doc.paragraphs[0]
    assert "Head Strategic Business Development" in p.text
    assert "heading" in p.style.name.lower()


def test_strong_becomes_bold_run():
    doc = _render_semantic_doc("<p>Delivery with <strong>vendor governance</strong> focus.</p>")

    p = doc.paragraphs[0]
    bold_runs = [r for r in p.runs if r.bold]
    assert bold_runs
    assert any("vendor governance" in r.text for r in bold_runs)


def test_em_becomes_italic_run():
    doc = _render_semantic_doc("<p>Strong fit for <em>insurance transformation</em> context.</p>")

    p = doc.paragraphs[0]
    italic_runs = [r for r in p.runs if r.italic]
    assert italic_runs
    assert any("insurance transformation" in r.text for r in italic_runs)


def test_li_becomes_list_bullet():
    html = "<ul><li>AWS</li><li>Azure</li><li>GCP</li></ul>"
    doc = _render_semantic_doc(html)

    bullet_paras = [p for p in doc.paragraphs if "list bullet" in p.style.name.lower()]
    assert len(bullet_paras) >= 3
    assert any("AWS" in p.text for p in bullet_paras)


def test_p_becomes_normal_para():
    doc = _render_semantic_doc("<p>Program leadership and cloud migration delivery.</p>")

    p = doc.paragraphs[0]
    assert "Program leadership" in p.text
    assert p.style.name.lower() in {"normal", "body text"} or "normal" in p.style.name.lower()


def test_nested_tags_preserved():
    html = "<p><strong>Bold</strong> and <em>italic</em> in one sentence.</p>"
    doc = _render_semantic_doc(html)

    p = doc.paragraphs[0]
    assert any(r.bold and "Bold" in r.text for r in p.runs)
    assert any(r.italic and "italic" in r.text for r in p.runs)


def test_full_cv_html_to_docx_round_trip():
    from app.services.premium_export import generate_premium_docx

    html = (
        "<h1>Manoel Benicio</h1>"
        "<h2>Professional Summary</h2><p>Delivery director focused on modernization and governance.</p>"
        "<h2>Experience</h2><h3>Head Strategic Business Development</h3><ul><li>Led cloud programs</li></ul>"
        "<h2>Education</h2><p>MBA - Solutions Architecture</p>"
        "<h2>Certifications</h2><ul><li>AWS Solutions Architect</li></ul>"
    )

    docx_bytes = generate_premium_docx(
        enhanced_cv_text=html,
        job_title="FBS Delivery Director",
        company="Capgemini",
        candidate_name="Manoel Benicio",
        skills_matched=["AWS", "Azure", "Leadership"],
        skills_missing=["Insurance domain"],
    )

    doc = Document(io.BytesIO(docx_bytes))
    text = _all_doc_text(doc)

    assert "MANOEL BENICIO" in text.upper()
    assert "PROFESSIONAL SUMMARY" in text.upper() or "Professional Summary" in text
    assert "EXPERIENCE" in text.upper() or "Experience" in text
    assert "AWS Solutions Architect" in text


def test_docx_ats_compliance():
    from app.services.premium_export import generate_premium_docx

    html = (
        "<h1>Manoel Benicio</h1>"
        "<h2>Summary</h2><p>Cloud modernization leader.</p>"
        "<h2>Experience</h2><ul><li>Program delivery</li><li>Governance</li></ul>"
    )

    docx_bytes = generate_premium_docx(
        enhanced_cv_text=html,
        job_title="FBS Delivery Director",
        company="Capgemini",
        candidate_name="Manoel Benicio",
        skills_matched=["Cloud"],
        skills_missing=[],
    )

    doc = Document(io.BytesIO(docx_bytes))
    assert len(doc.tables) == 0
    assert len(doc.inline_shapes) == 0

    # Single-column assertion (ATS-friendly): explicit multi-column layouts must not exist.
    for section in doc.sections:
        cols = section._sectPr.xpath("./w:cols")
        if cols:
            num = cols[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num")
            assert num in (None, "1")


# Module F: Backward Compatibility (F1-F4)

def test_generate_docx_still_accepts_plain_text():
    from app.services.premium_export import generate_premium_docx

    text = (
        "PROFESSIONAL SUMMARY\n"
        "Delivery leader in modernization.\n\n"
        "EXPERIENCE\n"
        "Head Strategic Business Development\n"
        "- Led enterprise programs"
    )

    docx_bytes = generate_premium_docx(
        enhanced_cv_text=text,
        job_title="Director",
        company="Capgemini",
        candidate_name="Manoel Benicio",
        skills_matched=["Delivery"],
        skills_missing=[],
    )

    doc = Document(io.BytesIO(docx_bytes))
    assert "Delivery leader" in _all_doc_text(doc)


def test_generate_docx_truncated_json_payload():
    from app.services.premium_export import generate_premium_docx

    payload = (
        '{"enhanced_cv":"<h1>Manoel Benicio</h1><p>Sao Paulo</p>'
        "<h2>Professional Summary</h2><p>Strong candidate"
    )

    docx_bytes = generate_premium_docx(
        enhanced_cv_text=payload,
        job_title="Director",
        company="Capgemini",
        candidate_name=None,
        skills_matched=["Cloud"],
        skills_missing=[],
    )

    doc = Document(io.BytesIO(docx_bytes))
    text = _all_doc_text(doc)
    assert '{"enhanced_cv"' not in text
    assert "Manoel Benicio" in text


def test_generate_docx_empty_input():
    from app.services.premium_export import generate_premium_docx

    docx_bytes = generate_premium_docx(
        enhanced_cv_text="",
        job_title="Director",
        company="Capgemini",
        candidate_name="Default Candidate",
        skills_matched=[],
        skills_missing=[],
    )

    doc = Document(io.BytesIO(docx_bytes))
    text = _all_doc_text(doc)
    assert "DEFAULT CANDIDATE" in text.upper()
    assert "Targeting:" in text


def test_existing_cv_versions_render_correctly():
    from app.services.premium_export import generate_premium_docx

    stored_enhanced_content = (
        '{"enhanced_cv":"<h1>Manoel Benicio</h1><p>Sao Paulo, Brazil | '
        'manoel.benicio@icloud.com</p><h2>Professional Summary</h2>'
        '<p>FBS Delivery Director with enterprise modernization background.</p>'
        '<h2>Experience</h2><h3>Head Strategic Business Development</h3>'
        '<ul><li>Managed budget and delivery transformation</li></ul>"}'
    )

    docx_bytes = generate_premium_docx(
        enhanced_cv_text=stored_enhanced_content,
        job_title="FBS Delivery Director",
        company="Capgemini",
        candidate_name="Manoel Benicio",
        skills_matched=["Cloud", "Governance"],
        skills_missing=["Insurance depth"],
    )

    doc = Document(io.BytesIO(docx_bytes))
    text = _all_doc_text(doc)

    assert "MANOEL BENICIO" in text.upper()
    assert "FBS Delivery Director" in text
    assert not re.search(r"\{\s*\"enhanced_cv\"", text)
