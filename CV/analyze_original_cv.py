"""Analyze the formatting of the original CV - save to file."""
from docx import Document

doc = Document(r"D:\VMs\Projetos\RFP_Automation_VF\AI_Job_Matcher\CV\ManoelBenicio_CV 2026.docx")

lines = []

# Page setup
for s in doc.sections:
    lines.append("=== PAGE SETUP ===")
    lines.append(f"Top margin: {s.top_margin.inches:.2f} in")
    lines.append(f"Bottom margin: {s.bottom_margin.inches:.2f} in")
    lines.append(f"Left margin: {s.left_margin.inches:.2f} in")
    lines.append(f"Right margin: {s.right_margin.inches:.2f} in")
    lines.append(f"Page width: {s.page_width.inches:.2f} in (A4)" if s.page_width.inches < 8.5 else f"Page width: {s.page_width.inches:.2f} in (Letter)")
    lines.append(f"Page height: {s.page_height.inches:.2f} in")

lines.append(f"\n=== PARAGRAPHS ({len(doc.paragraphs)} total) ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text[:90] if p.text else "[EMPTY]"
    style = p.style.name if p.style else "None"
    alignment = str(p.alignment) if p.alignment else "None/LEFT"
    lines.append(f"\nP{i:02d} | Style: {style} | Align: {alignment}")
    lines.append(f"     Text: {text}")
    for j, run in enumerate(p.runs):
        font = run.font
        name = font.name or "inherited"
        size = f"{font.size.pt}pt" if font.size else "inherited"
        bold = "BOLD" if font.bold else ""
        italic = "ITALIC" if font.italic else ""
        underline = "UNDERLINE" if font.underline else ""
        color = str(font.color.rgb) if font.color and font.color.rgb else "inherited"
        flags = " ".join(filter(None, [bold, italic, underline]))
        lines.append(f"     Run{j}: font={name} size={size} {flags} color={color}")
        lines.append(f"            \"{run.text[:70]}\"")

with open(r"D:\VMs\Projetos\RFP_Automation_VF\AI_Job_Matcher\CV\format_analysis.txt", "w", encoding="utf-8") as f:
    f.write("\n".join(lines))

print(f"Analysis saved. {len(doc.paragraphs)} paragraphs analyzed.")
