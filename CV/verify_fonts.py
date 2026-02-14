"""Verify font standardization in the generated DOCX."""
from docx import Document

doc = Document(r"D:\VMs\Projetos\RFP_Automation_VF\AI_Job_Matcher\CV\ManoelBenicio_CV_DeliveryDirector_2026.docx")

print("=== FONT STANDARDIZATION CHECK ===\n")

# Collect all unique font configs
configs = {}
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        f = run.font
        name = f.name or "MISSING"
        size = f"{f.size.pt}pt" if f.size else "MISSING"
        bold = "Bold" if f.bold else ""
        italic = "Italic" if f.italic else ""
        style = " ".join(filter(None, [bold, italic])) or "Regular"
        key = f"{name} {size} {style}"
        if key not in configs:
            configs[key] = {"count": 0, "example": ""}
        configs[key]["count"] += 1
        if not configs[key]["example"]:
            configs[key]["example"] = run.text[:50]

print(f"{'Font Config':<35} {'Count':>5}  Example")
print("-" * 90)
for key, val in sorted(configs.items()):
    print(f"{key:<35} {val['count']:>5}  \"{val['example']}\"")

# Check for any non-Cambria fonts
print("\n=== NON-CAMBRIA FONTS ===")
found_issue = False
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if run.font.name and run.font.name != "Cambria":
            print(f"  WARNING P{i}: font={run.font.name} text=\"{run.text[:40]}\"")
            found_issue = True
        if run.text.strip() and not run.font.size:
            print(f"  WARNING P{i}: size=MISSING text=\"{run.text[:40]}\"")
            found_issue = True

if not found_issue:
    print("  ✅ All fonts are Cambria with explicit sizes!")
