"""
Executive CV Generator — Manoel Benicio Filho
FBS Delivery Director @ Capgemini (Insurance Client)

Generates a professionally formatted DOCX CV matching the original
"ManoelBenicio_CV 2026.docx" style:
  - Font: Cambria throughout
  - Name: 24pt Bold
  - Section Headers: 14pt Bold
  - Job Titles: Bold (Company | Title format)
  - Dates: 10pt Italic
  - Body: 10pt regular
  - Page: A4, 1-inch margins all around
  - Alignment: Left
  - Colors: Black (no colors)
  - No borders, no decorative elements

Usage: python generate_executive_cv.py
Output: ManoelBenicio_CV_DeliveryDirector_2026.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ─────────────────────────────────────────────────────────
# CONFIGURATION
# ─────────────────────────────────────────────────────────
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "ManoelBenicio_CV_DeliveryDirector_2026.docx")

FONT = "Cambria"


def add_run(paragraph, text, size=10, bold=False, italic=False):
    """Add a styled run to a paragraph using Cambria font."""
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    # Set Cambria for complex scripts too
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>')
    rPr.append(rFonts)
    return run


def set_spacing(paragraph, before=0, after=0):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def add_empty_line(doc):
    """Add an empty paragraph as spacer."""
    p = doc.add_paragraph()
    set_spacing(p, 0, 0)
    add_run(p, "", size=10)
    return p


def add_section_header(doc, text):
    """Add a section header: Cambria 14pt Bold."""
    p = doc.add_paragraph()
    set_spacing(p, before=4, after=4)
    add_run(p, text, size=14, bold=True)
    return p


def add_job_header(doc, company, title):
    """Add a job header: Cambria Bold — 'Company | Title' format."""
    p = doc.add_paragraph()
    set_spacing(p, before=6, after=0)
    add_run(p, f"{company} | {title}", bold=True)
    return p


def add_dates(doc, location, dates):
    """Add location + dates: Cambria 10pt Italic."""
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=2)
    add_run(p, f"{location} — {dates}", size=10, italic=True)
    return p


def add_bullet(doc, text):
    """Add a body text paragraph (bullet content without bullet char)."""
    p = doc.add_paragraph()
    set_spacing(p, before=1, after=1)
    add_run(p, text, size=10)
    return p


# ─────────────────────────────────────────────────────────
# BUILD DOCUMENT
# ─────────────────────────────────────────────────────────
def generate_cv():
    doc = Document()

    # ── A4 page, 1-inch margins ──
    for section in doc.sections:
        section.page_width = Cm(21.0)   # A4 width
        section.page_height = Cm(29.7)  # A4 height
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ═══════════════════════════════════════════════════════
    # NAME
    # ═══════════════════════════════════════════════════════
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=0)
    add_run(p, "MANOEL BENICIO", size=24, bold=True)

    # ── Subtitle ──
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=0)
    add_run(p, "Technology Delivery Director | Insurance & Financial Services IT Programs | Digital Transformation", size=10)

    # ── Contact ──
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=0)
    add_run(p, (
        "São Paulo, Brazil\n"
        "Email: manoel.benicio@icloud.com\n"
        "Phone: +55-11-99364-4444\n"
        "LinkedIn: linkedin.com/in/manoel-benicio-filho\n"
        "Languages: English (Fluent) | Portuguese (Native) | Spanish (Fluent)"
    ), size=10)

    add_empty_line(doc)

    # ═══════════════════════════════════════════════════════
    # PROFILE
    # ═══════════════════════════════════════════════════════
    add_section_header(doc, "Profile")

    add_bullet(doc, (
        "Results-driven Delivery Director with 20+ years leading enterprise-scale IT programs "
        "across insurance, banking, financial services, telecommunications, and public safety. "
        "Currently Head of the Financial Services & Insurance (SSFF) segment at Indra/Minsait, "
        "doubling premium client revenue year-over-year across Santander, Mapfre, BrasilSeg, "
        "and Getnet. Previously managed Telefónica's largest BPO operation — the $100M/year "
        "'Espacio Mapfre' program — overseeing ~200 direct reports across Data, SAP, "
        "Infrastructure, and Cybersecurity practices, plus 2,000+ indirect resources across "
        "Brazil, Miami (USA), and Spain. Delivered end-to-end IT operations across insurance "
        "platforms including Tronweb (policy administration), SAP, and Dynatrace-monitored "
        "cloud workloads. Certified multi-cloud architect (AWS, Azure, GCP, OCI) and IT "
        "governance specialist (ITIL, ISO 20000, ISO 27001). Proven ability to translate "
        "C-suite strategy into actionable IT roadmaps, drive organizational change management "
        "at scale, and deliver measurable outcomes: 37.4% IT cost reduction, 50.5% customer "
        "revenue growth, and 65% infrastructure performance improvement."
    ))

    add_empty_line(doc)

    # ═══════════════════════════════════════════════════════
    # CORE COMPETENCIES
    # ═══════════════════════════════════════════════════════
    add_section_header(doc, "Core Competencies")

    add_bullet(doc, (
        "IT Strategy & Business Alignment • Insurance & Financial Services IT • Cloud Modernization & Migration • "
        "Solution Architecture & Design • Program & Portfolio Management • Vendor & Outsourcing Management • "
        "IT Governance (ITIL, COBIT, ISO 20000, ISO 27001) • Change Management & Organizational Transformation • "
        "P&L Oversight • Budget Ownership ($200M+) • Revenue Growth & Cost Optimization • CAPEX/OPEX Management • "
        "Insurance Platforms (Tronweb, SAP) • Policy Administration • Claims Processing • Underwriting Systems • "
        "Stakeholder & C-Suite Advisory • Cross-Functional Team Leadership (~200 direct / 2,000+ indirect) • "
        "Application Monitoring (Dynatrace) • Cybersecurity Operations • Network Infrastructure • Regulatory Compliance • "
        "Multi-Cloud Architecture (AWS, Azure, GCP, OCI) • Digital Transformation • Agile / DevOps / SDLC • "
        "Risk Assessment & Mitigation • Service Delivery (SLA/KPI) • Data Governance & Business Intelligence • "
        "Multi-Country Operations (Americas, EMEA) • Global Talent Management • BPO Operations Management"
    ))

    add_empty_line(doc)

    # ═══════════════════════════════════════════════════════
    # CERTIFICATIONS
    # ═══════════════════════════════════════════════════════
    add_section_header(doc, "Certifications")

    add_bullet(doc, (
        "ITIL Service Management Foundation • "
        "EXIN ISO/IEC 20000 Foundation (IT Service Management Standard) • "
        "ISO/IEC 27001 (Information Security Management) • "
        "AWS Certified Solutions Architect – Associate • "
        "AWS Certified Security – Specialty • "
        "Azure Solutions Architect (AZ-303) • "
        "Azure Cybersecurity Architect • "
        "Azure Security Engineer Associate • "
        "Azure Network Engineer Associate • "
        "Azure Database Administrator • "
        "Oracle Multi-Cloud Architect (OCI) • "
        "Google Associate Cloud Engineer"
    ))

    add_empty_line(doc)

    # ═══════════════════════════════════════════════════════
    # WORK EXPERIENCE
    # ═══════════════════════════════════════════════════════
    add_section_header(doc, "Work Experience")

    # ─── INDRA / MINSAIT — SSFF Head ───
    add_job_header(doc,
        "Indra Group / Minsait",
        "Head of Financial Services & Insurance (SSFF) Segment — Apps & Cloud Modernization"
    )
    add_dates(doc, "São Paulo, BR", "2023 - Present")

    bullets_indra1 = [
        "Lead the SSFF segment covering all private banks and insurers, including Santander, Mapfre, BrasilSeg, Getnet, and other premium clients — doubling segment headcount and revenue year-over-year for the past 2 consecutive years",
        "Manage enterprise delivery programs across Americas and EMEA, leading a diverse engineering organization of 50+ professionals including developers, architects, security, and infrastructure specialists",
        "Drove customer revenue growth from 25% to 50.5% over 2 years through innovative cloud modernization products and solutions tailored to insurance and banking client outcomes",
        "Managed annual budget of $200M, achieving 37.4% reduction in IT expenditure through strategic vendor consolidation, cloud cost optimization, and process re-engineering",
        "Spearheaded cloud modernization initiatives for insurance and banking clients, migrating legacy on-premise systems (policy administration, claims, core banking) to hybrid/multi-cloud architectures (AWS/Azure)",
        "Served as strategic advisor to C-suite stakeholders, presenting business cases, IT investment roadmaps, and transformation strategies to executive committees and board members at Santander, Mapfre, and BrasilSeg",
        "Championed IT governance and compliance standards, implementing architecture governance frameworks aligned with ITIL, COBIT, ISO 20000, and ISO 27001 best practices",
        "Led organizational change management at scale, driving cultural transformation through Indra's 'Líderes Transformadores' executive leadership program (2025)",
    ]
    for b in bullets_indra1:
        add_bullet(doc, b)

    add_empty_line(doc)

    # ─── INDRA / MINSAIT — Data & Analytics ───
    add_job_header(doc,
        "Indra Group / Minsait",
        "Sr Data & Analytics Practice Manager — Insurance & Health Services"
    )
    add_dates(doc, "São Paulo, BR", "2023 - Present")

    bullets_indra2 = [
        "Oversee large-scale IT programs in the insurance and health services sector, ensuring timely execution, quality assurance, and regulatory compliance for clients including Mapfre and BrasilSeg",
        "Led cross-functional teams of data engineers, architects, and analysts delivering data-driven strategies that inform underwriting decisions, claims processing optimization, and operational risk assessment",
        "Reduced project completion time by 25% through agile methodologies, DevOps automation, and improved sprint planning processes",
        "Implemented enterprise data governance frameworks, reducing data errors by 25% and improving data quality for insurance regulatory reporting and business intelligence",
        "Mentor and develop 15+ technical professionals through AWS, Azure, and GCP certification journeys, building a culture of continuous learning since 2024",
    ]
    for b in bullets_indra2:
        add_bullet(doc, b)

    add_empty_line(doc)

    # ─── ANDELA ───
    add_job_header(doc, "Andela", "Head Cloud & Data Professional Services")
    add_dates(doc, "New York, USA", "2022 - 2023")

    bullets_andela = [
        "Led 30+ developers, security, and infrastructure professionals for cloud and data program delivery across US and international clients",
        "Managed globally distributed engineering teams spanning India, Nigeria, Egypt, Europe, and the Americas — gaining deep expertise in cross-cultural leadership and remote program delivery across multiple time zones",
        "Facilitated successful legacy migrations and upgrades across cloud/data platforms, ensuring alignment with AWS/Azure Well-Architected Frameworks",
        "Built a high-performance culture rooted in ownership, inclusiveness, accountability, and urgency — delivering a 35% improvement in team velocity",
        "Managed vendor and outsourcing partner relationships, including SLA governance, performance reviews, and contract optimization across multiple geographies",
        "Drove organizational change management for technology adoption, coaching teams and stakeholders through cloud-first transformations",
    ]
    for b in bullets_andela:
        add_bullet(doc, b)

    add_empty_line(doc)

    # ─── TELEFÓNICA — Mapfre BPO ───
    add_job_header(doc,
        "Telefónica Tech",
        "Sr Cloud Operations Manager — Espacio Mapfre BPO Program ($100M/yr)"
    )
    add_dates(doc, "São Paulo, BR | Miami, USA | Madrid, Spain", "2021 - 2022")

    bullets_tef1 = [
        "Managed Telefónica Tech's largest BPO operation — the 'Espacio Mapfre' program ($100M+/year budget) — overseeing ~200 direct professionals across Data, SAP, Infrastructure, and Cybersecurity practices, plus 2,000+ indirect resources across Brazil, Miami (USA), and Spain",
        "Delivered end-to-end IT operations for Mapfre's insurance platforms, including Tronweb (policy administration system), SAP (enterprise resource planning), claims processing, underwriting systems, and customer-facing digital channels",
        "Managed application performance and infrastructure observability using Dynatrace for real-time monitoring of insurance workloads, ensuring 99.9% uptime for mission-critical policy and claims systems",
        "Reported directly to the COO, serving as a trusted strategic advisor on transformation initiatives — from business case development through execution and value realization",
        "Optimized mission-critical insurance batch processing workflows, migrating legacy batch jobs to cloud-native services on AWS — significantly reducing processing times, compute resource consumption, and operational costs for policy administration and claims processing cycles",
        "Orchestrated cloud infrastructure migration for Mapfre's insurance workloads, delivering a 30% cost decrease and 65% performance improvement through architecture optimization, workload right-sizing, and disaster recovery enhancements",
        "Led organizational change management to drive ITSM/ITIL adoption across BPO operations teams at Mapfre, Telefónica, and Claro — overcoming resistance from legacy operational cultures by demonstrating compliance with ITIL service catalog governance, embedding structured incident/change/problem management processes, and aligning operations with ISO 20000 standards",
        "Ensured regulatory compliance for insurance data handling, including data residency, privacy (LGPD/GDPR), business continuity, and information security requirements aligned with ISO 27001 standards",
        "Coordinated multi-country operations across 3 geographies (São Paulo, Miami, Madrid), managing cultural diversity, time zone challenges, and local regulatory requirements simultaneously",
    ]
    for b in bullets_tef1:
        add_bullet(doc, b)

    add_empty_line(doc)

    # ─── TELEFÓNICA — Mapfre Contracts ───
    add_job_header(doc,
        "Telefónica Tech",
        "Head Contracts for Cloud Services — Mapfre Insurance Account"
    )
    add_dates(doc, "São Paulo, BR | Miami, USA", "2020 - 2021")

    bullets_tef2 = [
        "Led major B2B contracts for the LATAM region at Mapfre (global insurance leader), managing the full lifecycle of cloud service agreements, SLAs, and compliance requirements for insurance technology platforms",
        "Partnered with Mapfre's insurance operations teams to migrate on-premise policy administration, claims processing, and underwriting systems to public cloud platforms (AWS/Azure)",
        "Managed datacenter teams in Brazil and Miami, overseeing CAPEX/OPEX budgets, capacity planning, and disaster recovery for Mapfre's mission-critical insurance workloads",
        "Delivered IT solutions aligned to insurance regulatory requirements, including data residency, privacy (LGPD/GDPR), and business continuity standards for multi-country operations",
        "Managed ASP and ISP vendor relationships, ensuring service level compliance and operational excellence across infrastructure and application tiers",
    ]
    for b in bullets_tef2:
        add_bullet(doc, b)

    add_empty_line(doc)

    # ─── NICE ───
    add_job_header(doc, "NICE", "Program Manager - Public Safety")
    add_dates(doc, "Dallas, USA", "2016 - 2020")

    bullets_nice = [
        "Provided leadership for developers, security, and infrastructure professionals in public safety program delivery across multiple US regions",
        "Main sponsor managing Business Partner contracts (SLAs, performance, training, delivery) across US regions delivering professional services",
        "Managed partners across US regions delivering professional services and enterprise system integrations",
        "Collaborated with engineering and product teams to design and deploy NICE solutions ensuring on-time, on-budget, and on-scope delivery with structured IT governance processes",
        "Handled executive-level escalations, resolving complex challenges with internal stakeholders and external partners through structured problem-solving and negotiation",
    ]
    for b in bullets_nice:
        add_bullet(doc, b)

    add_empty_line(doc)

    # ═══════════════════════════════════════════════════════
    # KEY ACHIEVEMENTS
    # ═══════════════════════════════════════════════════════
    add_section_header(doc, "Key Achievements & Metrics")

    achievements = [
        "$100M+ BPO Program: Managed Telefónica's largest BPO engagement (Espacio Mapfre) with $100M/yr budget across 3 countries",
        "Revenue Growth: Grew customer revenue from 25% to 50.5% at Indra/Minsait through cloud modernization solutions for insurance and banking",
        "SSFF Segment Growth: Doubled segment headcount and revenue year-over-year for 2 consecutive years across Santander, Mapfre, BrasilSeg, Getnet",
        "Cost Optimization: Achieved 37.4% reduction in IT expenditure on a $200M annual budget through strategic vendor consolidation",
        "Team Scale: Led ~200 direct professionals (Data, SAP, Infra, Cyber) and coordinated 2,000+ indirect resources across Americas and EMEA",
        "Performance Improvement: Delivered 65% cloud infrastructure performance gain for Mapfre's insurance workloads",
        "Batch Processing: Optimized insurance batch processing (policy admin, claims) by migrating to cloud-native services — reducing processing time and resource costs",
        "ITSM Transformation: Drove ITIL/ITSM adoption across resistant BPO operations at Mapfre, Telefónica, and Claro — embedding structured IT service management culture",
        "Delivery Acceleration: Reduced project completion time by 25% via agile methodologies and DevOps automation",
        "Data Quality: Reduced data errors by 25% through enterprise data governance framework implementation",
        "Team Velocity: Improved engineering team velocity by 35% through high-performance ownership-driven culture",
    ]
    for a in achievements:
        add_bullet(doc, a)

    add_empty_line(doc)

    # ═══════════════════════════════════════════════════════
    # INDUSTRY EXPERTISE
    # ═══════════════════════════════════════════════════════
    add_section_header(doc, "Industry Expertise")

    industries = [
        "Insurance (Primary): End-to-end BPO operations for Mapfre ($100M/yr) using Tronweb (policy admin), SAP, and Dynatrace; claims processing, underwriting systems, insurance regulatory compliance (LGPD/GDPR), data governance for insurance reporting, SSFF segment leadership covering Mapfre, BrasilSeg, and other insurers",
        "Banking & Financial Services: Private banking segment (Santander, Getnet), core banking modernization, financial regulatory compliance, data analytics for risk assessment",
        "Telecommunications & Media: B2B enterprise contracts, cloud migration for telco workloads, managed services, network operations",
        "Public Safety & Government: Mission-critical systems, multi-vendor program delivery, enterprise system integrations (NICE, Dallas USA)",
        "Health Services: Large-scale IT programs, data-driven strategies, regulatory compliance",
    ]
    for ind in industries:
        add_bullet(doc, ind)

    add_empty_line(doc)

    # ═══════════════════════════════════════════════════════
    # GLOBAL EXPERIENCE
    # ═══════════════════════════════════════════════════════
    add_section_header(doc, "Global Experience & Multi-Country Operations")

    add_bullet(doc, (
        "Extensive experience leading multi-country IT operations and globally distributed teams. "
        "At Telefónica Tech, managed the Espacio Mapfre program across Brazil (São Paulo), "
        "USA (Miami), and Spain (Madrid) simultaneously. At Andela, led engineering talent "
        "from India, Nigeria, Egypt, Europe, and the Americas. Fully bilingual in English and "
        "Portuguese with professional fluency in Spanish. Adept at navigating cultural "
        "differences, time zone challenges, and local regulatory requirements across "
        "Americas and EMEA regions."
    ))

    add_empty_line(doc)

    # ═══════════════════════════════════════════════════════
    # EDUCATION
    # ═══════════════════════════════════════════════════════
    add_section_header(doc, "Education")

    add_bullet(doc, "Master of Business Administration (MBA) - Solutions Architecture | FIAP | 2020")
    add_bullet(doc, "Bachelor's Degree - Computer Systems Networking and Telecommunications | 2017")

    add_empty_line(doc)

    # ═══════════════════════════════════════════════════════
    # THOUGHT LEADERSHIP
    # ═══════════════════════════════════════════════════════
    add_section_header(doc, "Thought Leadership & Professional Development")

    thought = [
        "Completed Indra Group's 'Líderes Transformadores' program (2025) — advanced executive leadership, strategic thinking, and cultural transformation",
        "Since 2024, actively mentoring Minsait professionals through AWS, Azure, and GCP certification journeys — building next-generation cloud talent",
        "Active contributor to AI strategy discussions; attendee at Microsoft AI Tour São Paulo (2026)",
        "Regular contributor on LinkedIn with insights on IT leadership, cloud architecture, insurance technology, and digital transformation strategy",
    ]
    for t in thought:
        add_bullet(doc, t)

    add_empty_line(doc)

    # ═══════════════════════════════════════════════════════
    # CAREER GOAL
    # ═══════════════════════════════════════════════════════
    add_section_header(doc, "Career Goal")

    add_bullet(doc, (
        "To lead enterprise technology delivery and transformation programs for a global insurance "
        "or financial services organization, leveraging 20+ years of hands-on experience in IT strategy, "
        "cloud modernization, and large-scale program management to drive measurable business outcomes, "
        "optimize operational efficiency, and deliver innovative technology solutions that create "
        "sustainable competitive advantage."
    ))

    # ═══════════════════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════════════════
    doc.save(OUTPUT_FILE)
    print(f"\n{'='*60}")
    print(f"  ✅ EXECUTIVE CV GENERATED SUCCESSFULLY")
    print(f"{'='*60}")
    print(f"  📄 File: {OUTPUT_FILE}")
    print(f"  📏 Format: DOCX (A4, 1-inch margins)")
    print(f"  🎨 Style: Cambria — matching original CV format")
    print(f"  📊 Target: Delivery Director — Insurance")
    print(f"  📋 Content: Enhanced with insurance platforms,")
    print(f"     batch processing, ITSM change mgmt")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    generate_cv()
