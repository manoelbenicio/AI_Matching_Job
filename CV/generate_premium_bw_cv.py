"""
Premium Fortune 500 Executive CV — Manoel Benicio Filho
═════════════════════════════════════════════════════════
TARGET ROLE: FBS Delivery Director @ Capgemini
TARGET CLIENT: Top-5 US P&C Insurer ($25B+ GWP, 10M+ households, 48K agents)

DESIGN: Modern C-Level Black & White
  - Font: Garamond — the Fortune 500 standard (annual reports, McKinsey, BCG)
  - Pure Black & White only
  - Thin black separator rules
  - Letter-spaced uppercase headers
  - Dash-indented bullets
  - 3-column competencies, 2-column certs
  - A4, tight margins

JD KEYWORD MAPPING (all explicitly addressed):
  ✓ Program/Project Management & Delivery
  ✓ Solution Architecture & IT Governance compliance
  ✓ IT Budget / Financial Modeling / Cost Centers ($200M+)
  ✓ Insurance Background (required)
  ✓ ASP, ISP, Data, Operations, Security retained staff
  ✓ Strategic IT Planning aligned with business priorities
  ✓ Stakeholder influence / C-Suite advisory
  ✓ Business change initiatives with technology component
  ✓ Translates business needs into attainable plans
  ✓ Vendor Management / Outsourcing Partners
  ✓ Leading Account Executives (Level I/II)
  ✓ Emotional Intelligence / Conflict Management
  ✓ Big Picture Thinking → business results
  ✓ Escalation handling (internal & external partners)
  ✓ Trend scouting / future customer needs
  ✓ Training, mentoring IT professionals
  ✓ Complex organizational structures / multiple BUs
  ✓ Delivery frameworks
  ✓ English Proficient ✓ / Master's preferred ✓

Usage: python generate_premium_bw_cv.py
Output: ManoelBenicio_CV_Premium_F500_2026.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "ManoelBenicio_CV_Premium_F500_2026.docx")

BLACK = RGBColor(0, 0, 0)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MID_GRAY = RGBColor(0x66, 0x66, 0x66)

FONT_HEADING = "Garamond"
FONT_BODY = "Garamond"


# ─────────────────────────────────────────────────────────
# FORMATTING HELPERS
# ─────────────────────────────────────────────────────────

def set_font(run, font_name, size, bold=False, italic=False, color=BLACK, caps=False, spacing=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
    if caps:
        rPr.append(parse_xml(f'<w:caps {nsdecls("w")} w:val="true"/>'))
    if spacing:
        rPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:val="{spacing}"/>'))
    return run


def set_spacing(p, before=0, after=0, line_spacing=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)


def add_thin_rule(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=2, after=2)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/></w:pBdr>'
    ))
    return p


def add_thick_rule(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=0)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/></w:pBdr>'
    ))
    return p


def add_spacer(doc, pts=4):
    p = doc.add_paragraph()
    set_spacing(p, 0, 0)
    run = p.add_run("")
    run.font.size = Pt(pts)
    return p


def add_section_header(doc, title):
    add_thin_rule(doc)
    p = doc.add_paragraph()
    set_spacing(p, before=4, after=4)
    run = p.add_run(title.upper())
    set_font(run, FONT_BODY, 11, bold=True, color=BLACK, spacing=40)
    return p


def add_job_block(doc, title, company, location_dates):
    p = doc.add_paragraph()
    set_spacing(p, before=8, after=0)
    run = p.add_run(title)
    set_font(run, FONT_BODY, 10.5, bold=True, color=BLACK)
    p2 = doc.add_paragraph()
    set_spacing(p2, before=0, after=2)
    run2 = p2.add_run(f"{company}  |  {location_dates}")
    set_font(run2, FONT_BODY, 9.5, italic=True, color=MID_GRAY)
    return p, p2


def add_bullet(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=1, after=1)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    run = p.add_run("–  " + text)
    set_font(run, FONT_BODY, 10, color=DARK_GRAY)
    return p


def add_body(doc, text, size=10, color=DARK_GRAY, bold=False):
    p = doc.add_paragraph()
    set_spacing(p, before=1, after=1)
    run = p.add_run(text)
    set_font(run, FONT_BODY, size, bold=bold, color=color)
    return p


def add_achievement_row(doc, label, value):
    p = doc.add_paragraph()
    set_spacing(p, before=2, after=2)
    p.paragraph_format.left_indent = Inches(0.25)
    run_l = p.add_run(label)
    set_font(run_l, FONT_BODY, 10, bold=True, color=BLACK)
    run_s = p.add_run("  ·····  ")
    set_font(run_s, FONT_BODY, 10, color=MID_GRAY)
    run_v = p.add_run(value)
    set_font(run_v, FONT_BODY, 10, color=DARK_GRAY)
    return p


def add_tab_row(doc, items, tab_positions):
    """Add a tab-aligned row."""
    p = doc.add_paragraph()
    set_spacing(p, before=1, after=1)
    for i, item in enumerate(items):
        if i > 0:
            run_t = p.add_run("\t")
            set_font(run_t, FONT_BODY, 9)
        run = p.add_run(f"·  {item}")
        set_font(run, FONT_BODY, 9, color=DARK_GRAY)
    pPr = p._p.get_or_add_pPr()
    tabs_xml = f'<w:tabs {nsdecls("w")}>'
    for pos in tab_positions:
        tabs_xml += f'<w:tab w:val="left" w:pos="{pos}"/>'
    tabs_xml += '</w:tabs>'
    pPr.append(parse_xml(tabs_xml))
    return p


def add_tab_header(doc, titles, tab_positions):
    """Add tab-aligned column headers."""
    p = doc.add_paragraph()
    set_spacing(p, before=2, after=4)
    for i, title in enumerate(titles):
        if i > 0:
            run_t = p.add_run("\t")
            set_font(run_t, FONT_BODY, 8.5)
        run = p.add_run(title)
        set_font(run, FONT_BODY, 8.5, bold=True, color=BLACK, spacing=20)
    pPr = p._p.get_or_add_pPr()
    tabs_xml = f'<w:tabs {nsdecls("w")}>'
    for pos in tab_positions:
        tabs_xml += f'<w:tab w:val="left" w:pos="{pos}"/>'
    tabs_xml += '</w:tabs>'
    pPr.append(parse_xml(tabs_xml))
    return p


# ═══════════════════════════════════════════════════════════
# BUILD THE PREMIUM CV
# ═══════════════════════════════════════════════════════════
def generate_premium_cv():
    doc = Document()

    # ── A4, tight margins ──
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(10)
    style.font.color.rgb = DARK_GRAY

    # ════════════════════════════════════════════════════════
    # NAME BLOCK
    # ════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=0)
    run = p.add_run("MANOEL BENICIO FILHO")
    set_font(run, FONT_HEADING, 28, bold=False, color=BLACK, spacing=80)

    p = doc.add_paragraph()
    set_spacing(p, before=2, after=0)
    run = p.add_run("IT Delivery Director  |  Insurance & Financial Services Programs  |  Enterprise Modernization")
    set_font(run, FONT_HEADING, 10.5, color=MID_GRAY)

    add_spacer(doc, 2)

    p = doc.add_paragraph()
    set_spacing(p, before=1, after=0)
    run = p.add_run("São Paulo, Brazil  ·  manoel.benicio@icloud.com  ·  +55 11 99364-4444")
    set_font(run, FONT_BODY, 9, color=DARK_GRAY)

    p = doc.add_paragraph()
    set_spacing(p, before=1, after=0)
    run = p.add_run("linkedin.com/in/manoel-benicio-filho  ·  English (Fluent)  |  Portuguese (Native)  |  Spanish (Fluent)")
    set_font(run, FONT_BODY, 9, color=DARK_GRAY)

    add_thick_rule(doc)

    # ════════════════════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════
    add_section_header(doc, "Executive Summary")

    add_body(doc, (
        "IT Delivery Director with 20+ years leading enterprise-scale programs across "
        "insurance, financial services, telecommunications, and public safety sectors. "
        "Currently heading the Financial Services & Insurance (SSFF) segment at Indra/Minsait "
        "— doubling premium client revenue across Santander, Mapfre, BrasilSeg, and Getnet. "
        "Previously managed Telefónica's largest BPO operation — the $100M/year Espacio Mapfre "
        "insurance program — overseeing ~200 direct reports (Data, SAP, Infrastructure, "
        "Cybersecurity) and 2,000+ indirect resources across Brazil, Miami, and Spain. Proven "
        "track record managing $200M+ IT budgets, driving 37.4% cost reduction, and delivering "
        "50.5% customer revenue growth through solution architecture, cloud modernization, and "
        "IT governance. Certified multi-cloud architect (AWS, Azure, GCP, OCI) with deep "
        "expertise in IT governance (ITIL, COBIT, ISO 20000, ISO 27001), vendor/outsourcing "
        "management, stakeholder influence at C-suite level, and translating business strategy "
        "into actionable technology roadmaps."
    ))

    # ════════════════════════════════════════════════════════
    # CORE COMPETENCIES — 3-column (mirroring JD exactly)
    # ════════════════════════════════════════════════════════
    add_section_header(doc, "Core Competencies")

    tab_pos = [3200, 6400]

    add_tab_header(doc,
        ["STRATEGIC & BUSINESS", "TECHNICAL & DELIVERY", "GOVERNANCE & LEADERSHIP"],
        tab_pos
    )

    competencies = [
        ("IT Strategy & Business Alignment",  "Solution Architecture",              "IT Governance (ITIL/COBIT)"),
        ("Insurance & Financial Services IT",  "Program & Project Management",       "Compliance (ISO 20000/27001)"),
        ("IT Budget & Financial Modeling",     "Cloud Modernization & Migration",    "Vendor & Outsourcing Mgmt"),
        ("P&L / Cost Center Management",       "Application Modernization",          "Change Management"),
        ("Executive & Stakeholder Advisory",   "Delivery Framework Design",          "Cross-Functional Leadership"),
        ("Revenue & Cost Optimization",        "ASP / ISP / Data / Ops / Security",  "Training & Mentoring"),
        ("Big Picture Thinking",               "Insurance Platforms (Tronweb/SAP)",   "Escalation Management"),
        ("Trend Scouting & Innovation",        "Multi-Cloud (AWS/Azure/GCP/OCI)",     "Emotional Intelligence"),
    ]
    for row in competencies:
        add_tab_row(doc, row, tab_pos)

    # ════════════════════════════════════════════════════════
    # PROFESSIONAL EXPERIENCE
    # ════════════════════════════════════════════════════════
    add_section_header(doc, "Professional Experience")

    # ─── INDRA 1: SSFF Head ───
    add_job_block(doc,
        "Head of Financial Services & Insurance (SSFF) — Apps & Cloud Modernization",
        "Indra Group / Minsait",
        "São Paulo, Brazil  ·  2023 – Present"
    )
    for b in [
        "Lead the SSFF segment covering all private banks and insurers (Santander, Mapfre, BrasilSeg, Getnet) — doubled segment revenue YoY for 2 consecutive years across multiple business units",
        "Manage program delivery across Americas & EMEA, leading 50+ professionals (developers, solution architects, security, infrastructure) delivering business solutions of moderate to high complexity",
        "Own $200M annual IT budget across multiple cost centers; achieved 37.4% expenditure reduction through strategic vendor consolidation, financial modeling, and cloud cost optimization",
        "Drove customer revenue growth from 25% to 50.5% through solution architecture design and cloud modernization products tailored to insurance and banking client outcomes",
        "Translate C-suite business strategy into actionable IT plans and delivery frameworks — present investment roadmaps and transformation business cases to executive committees and board members",
        "Spearhead legacy-to-cloud migrations for insurance clients (policy administration, claims processing, core banking) ensuring alignment of IT strategy with business priorities",
        "Lead and mentor Account Executives (Level I/II), developing next-generation IT delivery leaders through structured training and certification programs",
        "Implement IT governance frameworks and optimize organizational processes aligned with ITIL, COBIT, ISO 20000, and ISO 27001 standards, ensuring compliance across all delivery programs",
        "Identify and pursue business and efficiency opportunities as trend scout — proactively identifying future customer needs and translating into technology investment proposals",
        "Completed Indra's 'Líderes Transformadores' executive leadership program (2025) — emotional intelligence, conflict management, empathy, and self-awareness development",
    ]:
        add_bullet(doc, b)

    # ─── INDRA 2: Data & Analytics ───
    add_job_block(doc,
        "Senior Data & Analytics Practice Manager — Insurance & Health Services",
        "Indra Group / Minsait",
        "São Paulo, Brazil  ·  2023 – Present"
    )
    for b in [
        "Oversee IT programs in insurance and health services verticals, ensuring timely execution, quality assurance, and regulatory compliance across complex organizational structures",
        "Lead cross-functional teams of data engineers and architects, delivering data-driven strategies for underwriting decisions, claims processing optimization, and operational risk assessment",
        "Reduced project completion time by 25% through agile methodologies, DevOps automation, and delivery framework optimization",
        "Implemented enterprise data governance framework, reducing data errors by 25% for insurance regulatory reporting and business intelligence",
        "Mentor 15+ IT professionals through AWS/Azure/GCP certification journeys, building a culture of continuous learning and technical excellence since 2024",
    ]:
        add_bullet(doc, b)

    # ─── ANDELA ───
    add_job_block(doc,
        "Head of Cloud & Data Professional Services",
        "Andela",
        "New York, USA  ·  2022 – 2023"
    )
    for b in [
        "Led 30+ developers, security, and infrastructure professionals in cloud and data program delivery across US and international clients",
        "Managed globally distributed engineering teams (India, Nigeria, Egypt, Europe, Americas) — deep expertise in cross-cultural leadership and remote delivery across time zones",
        "Facilitated legacy migrations and technology upgrades aligned to AWS/Azure Well-Architected Frameworks and IT governance best practices",
        "Built high-performance culture rooted in ownership, inclusiveness, and accountability — delivered 35% team velocity improvement",
        "Managed vendor relationships, outsourcing partner performance, SLA governance, and contract optimization across multiple geographies",
        "Drove organizational change management for cloud-first technology adoption, coaching teams and stakeholders through business change initiatives with significant technology components",
    ]:
        add_bullet(doc, b)

    # ─── TELEFÓNICA 1: Mapfre BPO ───
    add_job_block(doc,
        "Senior Cloud Operations Manager — Espacio Mapfre BPO Program ($100M/yr)",
        "Telefónica Tech",
        "São Paulo | Miami | Madrid  ·  2021 – 2022"
    )
    for b in [
        "Managed Telefónica Tech's largest BPO operation ($100M/yr) for Mapfre — one of the world's largest insurance groups — overseeing ~200 direct professionals across Data, SAP, Infrastructure, and Cybersecurity practices, plus 2,000+ indirect resources across Brazil, USA, and Spain",
        "Delivered end-to-end IT operations for Mapfre's insurance platforms: Tronweb (policy administration), SAP (ERP), claims processing, underwriting systems, and customer-facing digital channels",
        "Responsible for local ASP, ISP, Data, Operations, and Security retained staff — maintaining necessary knowledge capital and ensuring service continuity for mission-critical insurance workloads",
        "Reported directly to COO as trusted strategic advisor — translated strategic intent into actions, developed project concepts, created plans, and oversaw implementation of business solutions",
        "Handled customer challenges and escalations from internal and external business partners — used high emotional intelligence to manage conflict, show empathy, and drive resolution",
        "Orchestrated cloud infrastructure migration delivering 30% cost decrease and 65% performance improvement through solution architecture optimization and workload right-sizing",
        "Led organizational change management driving ITSM/ITIL adoption across BPO operations at Mapfre, Telefónica, and Claro — embedding structured governance processes aligned with ISO 20000",
        "Ensured regulatory compliance for insurance data handling: data residency, privacy (LGPD/GDPR), business continuity, and information security aligned with ISO 27001",
        "Coordinated multi-country operations across 3 geographies, navigating complex organizational structures, cultural diversity, and local regulatory requirements simultaneously",
    ]:
        add_bullet(doc, b)

    # ─── TELEFÓNICA 2: Mapfre Contracts ───
    add_job_block(doc,
        "Head of Contracts for Cloud Services — Mapfre Insurance Account",
        "Telefónica Tech",
        "São Paulo | Miami  ·  2020 – 2021"
    )
    for b in [
        "Led major B2B contracts for Mapfre across LATAM — managing full lifecycle of cloud service agreements, SLAs, financial modeling, and compliance for insurance technology platforms",
        "Partnered with Mapfre's insurance operations teams to migrate on-premise policy administration, claims, and underwriting systems to public cloud (AWS/Azure)",
        "Managed datacenter teams in Brazil and Miami — overseeing CAPEX/OPEX budgets, capacity planning, and disaster recovery for mission-critical insurance workloads",
        "Managed ASP and ISP vendor relationships, ensuring service level compliance and operational excellence across infrastructure and application tiers",
        "Ensured compliance with insurance regulatory requirements (LGPD/GDPR) for data handling across multi-country operations",
    ]:
        add_bullet(doc, b)

    # ─── NICE ───
    add_job_block(doc,
        "Program Manager — Public Safety",
        "NICE",
        "Dallas, USA  ·  2016 – 2020"
    )
    for b in [
        "Led multi-vendor program delivery of enterprise public safety solutions across US regions — managing developers, security, and infrastructure professionals",
        "Main sponsor managing Business Partner contracts: SLAs, performance management, training, and delivery across multiple US regions",
        "Designed delivery frameworks and oversaw implementation of enterprise-grade solutions on-time, on-budget, and on-scope with structured IT governance",
        "Handled executive-level escalations from internal and external partners — resolved complex challenges through structured problem-solving and stakeholder negotiation",
        "Identified business and efficiency opportunities, proactively scouting emerging technology trends to address future customer needs",
    ]:
        add_bullet(doc, b)

    # ════════════════════════════════════════════════════════
    # KEY ACHIEVEMENTS
    # ════════════════════════════════════════════════════════
    add_section_header(doc, "Key Achievements")

    for label, value in [
        ("Revenue Growth",          "25% → 50.5% customer revenue via cloud modernization solutions"),
        ("Cost Reduction",          "37.4% IT expenditure reduction on $200M+ annual budget"),
        ("BPO Program Scale",       "$100M/yr Mapfre program — ~200 direct + 2,000+ indirect resources"),
        ("Segment Growth",          "Doubled SSFF segment headcount and revenue YoY for 2 consecutive years"),
        ("Performance Gain",        "65% cloud infrastructure improvement + 30% cost decrease"),
        ("Delivery Acceleration",   "25% project time reduction via agile/DevOps / delivery frameworks"),
        ("Data Quality",            "25% error reduction through enterprise data governance"),
        ("Team Velocity",           "35% improvement through ownership-driven high-performance culture"),
    ]:
        add_achievement_row(doc, label, value)

    # ════════════════════════════════════════════════════════
    # CERTIFICATIONS — 2-column
    # ════════════════════════════════════════════════════════
    add_section_header(doc, "Certifications")

    cert_tab = [4800]
    add_tab_header(doc, ["CLOUD & ARCHITECTURE", "SECURITY & GOVERNANCE"], cert_tab)

    for left, right in [
        ("AWS Solutions Architect – Associate",    "AWS Security – Specialty"),
        ("Azure Solutions Architect (AZ-303)",     "Azure Cybersecurity Architect"),
        ("Oracle Multi-Cloud Architect (OCI)",     "Azure Security Engineer Associate"),
        ("Google Associate Cloud Engineer",        "ITIL Service Management Foundation"),
        ("Azure Network Engineer Associate",       "EXIN ISO/IEC 20000 Foundation"),
        ("Azure Database Administrator",           "ISO/IEC 27001 (Information Security)"),
    ]:
        p = doc.add_paragraph()
        set_spacing(p, before=1, after=1)
        run_l = p.add_run(f"·  {left}")
        set_font(run_l, FONT_BODY, 9, color=DARK_GRAY)
        run_t = p.add_run("\t")
        set_font(run_t, FONT_BODY, 9)
        run_r = p.add_run(f"·  {right}")
        set_font(run_r, FONT_BODY, 9, color=DARK_GRAY)
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="left" w:pos="4800"/></w:tabs>'))

    # ════════════════════════════════════════════════════════
    # EDUCATION
    # ════════════════════════════════════════════════════════
    add_section_header(doc, "Education")

    add_body(doc, "MBA — Solutions Architecture  |  FIAP  |  2020", bold=False)
    add_body(doc, "Bachelor's — Computer Systems, Networking & Telecommunications  |  2017", bold=False)

    # ════════════════════════════════════════════════════════
    # INDUSTRY EXPERTISE
    # ════════════════════════════════════════════════════════
    add_section_header(doc, "Industry Expertise")

    for label, desc in [
        ("Insurance (Primary)",          "End-to-end BPO operations ($100M/yr Mapfre), policy admin (Tronweb), SAP, claims, underwriting, Dynatrace monitoring, LGPD/GDPR compliance, SSFF segment leadership"),
        ("Banking & Financial Services", "Private banking (Santander, Getnet), core banking modernization, financial modeling, data analytics, regulatory compliance"),
        ("Telecommunications & Media",   "B2B enterprise contracts, cloud migration, managed services, network infrastructure operations"),
        ("Public Safety & Government",   "Mission-critical systems, multi-vendor program delivery, enterprise integrations (NICE, Dallas)"),
        ("Health Services",              "Large-scale IT programs, data-driven strategies, regulatory compliance"),
    ]:
        p = doc.add_paragraph()
        set_spacing(p, before=2, after=2)
        p.paragraph_format.left_indent = Inches(0.25)
        run_l = p.add_run(f"{label}  ")
        set_font(run_l, FONT_BODY, 10, bold=True, color=BLACK)
        run_s = p.add_run("—  ")
        set_font(run_s, FONT_BODY, 10, color=MID_GRAY)
        run_d = p.add_run(desc)
        set_font(run_d, FONT_BODY, 10, color=DARK_GRAY)

    add_thin_rule(doc)

    # ════════════════════════════════════════════════════════
    # SAVE
    # ════════════════════════════════════════════════════════
    doc.save(OUTPUT_FILE)
    print(f"\n{'═'*60}")
    print(f"  ✅ PREMIUM FORTUNE 500 CV GENERATED")
    print(f"{'═'*60}")
    print(f"  📄 {OUTPUT_FILE}")
    print(f"  🎯 Target: Capgemini FBS Delivery Director")
    print(f"     → Top-5 US P&C Insurer ($25B+ GWP)")
    print(f"  🎨 Design: Modern B&W C-Level")
    print(f"     · Garamond — Fortune 500 corporate serif")
    print(f"     · Pure Black & White")
    print(f"     · 3-col competencies, 2-col certs")
    print(f"  📋 JD Keywords: ALL mapped")
    print(f"     ✓ Insurance Background (primary)")
    print(f"     ✓ IT Governance / Solution Architecture")
    print(f"     ✓ ASP/ISP/Data/Ops/Security retained staff")
    print(f"     ✓ $200M+ Budget / Financial Modeling")
    print(f"     ✓ Vendor/Outsourcing Management")
    print(f"     ✓ Escalation Handling / EQ")
    print(f"     ✓ Trend Scouting / Big Picture Thinking")
    print(f"     ✓ Account Executive Leadership (I/II)")
    print(f"     ✓ Training & Mentoring")
    print(f"{'═'*60}\n")


if __name__ == "__main__":
    generate_premium_cv()
